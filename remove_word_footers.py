import os
import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def add_field(run, field_name):
    """
    插入Word域，例如：
    PAGE
    NUMPAGES
    """

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"

    fld_separate.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_separate)
    r.append(fld_end)


def add_page_footer(paragraph):

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = paragraph.add_run("第 ")

    run = paragraph.add_run()
    add_field(run, "PAGE")

    paragraph.add_run(" 页，共 ")

    run = paragraph.add_run()
    add_field(run, "NUMPAGES")

    paragraph.add_run(" 页")


def remove_footer_and_add_page(file_path):

    try:

        doc = Document(file_path)

        for section in doc.sections:

            footer = section.footer

            # 删除页脚所有内容
            footer_element = footer._element

            for child in list(footer_element):
                footer_element.remove(child)

            # 添加新的页脚
            paragraph = footer.add_paragraph()

            add_page_footer(paragraph)

        doc.save(file_path)

        print(f"√ 已处理：{file_path}")

    except Exception as e:

        print(f"× 处理失败：{file_path}")
        print(e)


def process_all_docx(root_folder):

    for root, dirs, files in os.walk(root_folder):

        for file in files:

            # 跳过Word临时文件
            if file.startswith("~$"):
                continue

            if file.lower().endswith(".docx"):

                full_path = os.path.join(root, file)

                remove_footer_and_add_page(full_path)


if __name__ == "__main__":

    if getattr(sys, "frozen", False):
        ROOT_FOLDER = os.path.dirname(sys.executable)
    else:
        ROOT_FOLDER = os.path.dirname(os.path.abspath(__file__))

    print("=" * 60)
    print("Word 页脚处理工具")
    print("=" * 60)
    print(f"处理目录：{ROOT_FOLDER}")
    print()

    process_all_docx(ROOT_FOLDER)

    print()
    print("=" * 60)
    print("全部处理完成！")
    print("=" * 60)

    input("按回车键退出...")